from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from typing import Dict, Any
import logging

def create_lendo_credit_file(summary_data: Dict[str, Any], output_filename="Lendo Credit File - ADK AGENT.docx"):
    """
    Creates a DOCX file mimicking the structure, content, and basic styles
    of the "Lendo Credit File - ADK AGENT.docx" file.

    Args:
        output_filename (str): The name of the DOCX file to create.
    """
    logging.basicConfig(
        level=logging.DEBUG,
        filename='adk_agent.log',  # This file will be created in the current working directory
        filemode='w',              # 'w' to overwrite each run; use 'a' to append
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    logger = logging.getLogger(__name__)

    logger.info(f"Summary data: {summary_data}")
    logger.info(f"output_filename: {output_filename}")

    # Start creating document file
    document = Document()

    # Set default font to Calibri (common default, might need to be adjusted if original is different)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11) # Default font size

    # Set very narrow margins in inches
    section = document.sections[0]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # --- Credit Decision ---
    document.add_heading('Credit Decision', level=1)

    # Risk Recommendation Table
    document.add_paragraph() # Add a little spacing
    table = document.add_table(rows=2, cols=3)
    table.style = 'Table Grid' # Apply a default grid style

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Risk Recommendation'
    hdr_cells[1].text = 'Approving Authority'
    hdr_cells[2].text = 'Underwriter Name'

    # Set blue background for header cells
    for cell in hdr_cells:
        #cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set background color to header cells
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    # Data Row
    data_cells = table.rows[1].cells
    data_cells[0].text = summary_data.get("finalDecision", "Not Recommend for financing (default)")
    #data_cells[0].text = 'Approved as Requested'
    data_cells[1].text = 'AI Credit Risk Officer'
    data_cells[2].text = 'Google ADK agent developed by Emmad, Imran, Saad, Shafeeque, Sumayyah, Hamza'
    for cell in data_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph()  # One blank line

    # --- Product Table ---
    document.add_paragraph() # Add a little spacing
    table = document.add_table(rows=3, cols=9)
    table.style = 'Table Grid'

    # Header Row 1
    hdr_cells = table.rows[0].cells
    # Merge cells[0] and cells[1]
    merged_cell = hdr_cells[0].merge(hdr_cells[1])
    merged_cell.text = 'Product'
    hdr_cells[2].text = 'Limit'
    hdr_cells[3].text = 'Available % For First Utilization'
    #hdr_cells[4].text = 'Neo Risk Rating'
    hdr_cells[4].text = 'Internal Risk Rating'
    #hdr_cells[6].text = 'Wiser Funding Risk Rating'
    hdr_cells[5].text = 'Pricing'
    hdr_cells[6].text = 'Mgmt Fee'
    hdr_cells[7].text = 'Tenor'
    hdr_cells[8].text = 'Advance Rate'
    for cell in hdr_cells:
        #cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Set background color to header cells
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    # Data Row 1 (A, Invoice Discounting)
    data_cells = table.rows[1].cells
    data_cells[0].text = 'A'
    data_cells[1].text = 'Invoice Discounting '
    data_cells[2].text = '3.5mn'
    data_cells[3].text = '100%'
    #data_cells[4].text = 'C'
    data_cells[4].text = f"{summary_data.get("riskRating", "N/A")} ({summary_data.get("simahScore", "N/A")})"
    #data_cells[6].text = 'B'
    data_cells[5].text = '18%'
    data_cells[6].text = '3%'
    data_cells[7].text = 'Based on Contracts ad repayment history'
    data_cells[8].text = '80%'
    for cell in data_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data Row 2 (Total)
    data_cells = table.rows[2].cells
    data_cells[0].text = 'Total'
    data_cells[1].text = '' # Merged cell originally, but `python-docx` handles width well
    data_cells[2].text = '3.5mn'
    data_cells[3].text = ''
    #data_cells[4].text = ''
    data_cells[4].text = ''
    #data_cells[6].text = ''
    data_cells[5].text = ''
    data_cells[6].text = ''
    data_cells[7].text = ''
    for cell in data_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Apply bold to 'Total'
    #data_cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()  # One blank line

    # --- Covenants/conditions and Security ---
    document.add_paragraph() # Spacing
    table = document.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # Row 1
    cell1 = table.rows[0].cells[0]
    cell2 = table.rows[0].cells[1]

    # Format left cell
    cell1.text = "Covenants/conditions"
    shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    cell1._tc.get_or_add_tcPr().append(shading)
    #cell1.paragraphs[0].runs[0].bold = True
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Format right cell
    p = cell2.paragraphs[0]
    p.add_run("Precedent Conditions")
    p = cell2.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Valid Zakat Certificated.")

    p = cell2.add_paragraph()
    p.add_run("Internal Conditions")
    p = cell2.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Max tenor is up to 180 days.")

    # Row 2
    cell1 = table.rows[1].cells[0]
    cell2 = table.rows[1].cells[1]

    cell1.text = "Security"
    shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    cell1._tc.get_or_add_tcPr().append(shading)
    #cell1.paragraphs[0].runs[0].bold = True
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Bullets for security items
    p = cell2.paragraphs[0]
    p.clear()  # Clear default paragraph
    p = cell2.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Order Note from Etihad Jeddah Factory company with the value of SAR 4,130,000/-")

    p = cell2.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Personal Order Note from Each Anwar Alhashri and Fahad Alhabrdi with the value of SAR 4,130,000 /-")

    document.add_paragraph()  # One blank line

    # --- Approved buyers Table ---
    document.add_paragraph()
    table = document.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Approved buyers'
    hdr_cells[1].text = 'CAP'
    hdr_cells[2].text = 'Tenor'
    for cell in hdr_cells:
        #cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set background color to header cells
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    # Data Rows
    data = [
        ['Govt & Semi-Govt', '100%', '--'],
        ['Well known Cooperation', '50%', '--'],
        ['Ajeej Steel Manufacturing Co', '50%', '--'],
        ['Abdulkarim Alrajhi Steel', '50%', '--'],
        ['Kema Gulg Trading Co', '50%', '--']
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_page_break()

    # --- Deal Deets Table ---
    document.add_paragraph()
    table = document.add_table(rows=9, cols=6) # Based on text parsing
    table.style = 'Table Grid'

    # Merge the first row's 6 cells to make a single title row
    header_row = table.rows[0]
    merged_cell = header_row.cells[0].merge(header_row.cells[1])
    for i in range(2, 6):
        merged_cell = merged_cell.merge(header_row.cells[i])

    # Set title text
    merged_cell.text = "Deal Deets"
    merged_paragraph = merged_cell.paragraphs[0]
    #merged_paragraph.runs[0].bold = True
    merged_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Set background color to header cells
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    merged_cell._tc.get_or_add_tcPr().append(shading_elm_1)

    # Headers for 'Deal Deets'
    deal_deets_headers = [
        'Credit File Date', '02-06-2025', 'Next Review Date', '31-06-2026', 'Last Review Date', 'New',
        'Entity Name', 'Etihad Jeddah Factory Co', 'BR#', '5637', 'CR# ', '4030298618',
        'Assessment', 'NewBiz / Full Review', 'Entity#', '7007994770', 'Legal Structure', 'LLC',
        'Type of Product', 'Invoice Finance', 'Incorporation Date ', '2018', 'Business Address ', 'Jeddah',
        'Industry / Sector', 'Manufacturing ', 'Zakat ', 'Not valid', 'Nitaqat ', 'Grean ',
        '# of branches', '3', 'RAM', 'High Risk', 'PEPs Sanctions', 'No',
        'Deal Source', 'RM', 'Relationship with Lendo (mos)', 'New', 'Watchlist Status', 'NewBiz'
    ]

    for i in range(7):
        cells = table.rows[i+1].cells
        for j in range(6):
            cell_text = deal_deets_headers[i * 6 + j]
            cells[j].text = cell_text
            # Bold first column items in each row
            if j % 2 == 0:
                #cells[j].paragraphs[0].runs[0].bold = True
                # Set background color to header cells
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
                cells[j]._tc.get_or_add_tcPr().append(shading_elm_1)
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # set text for last row first column
    cells = table.rows[8].cells
    cells[0].text = "Client Request"
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

    # Merge the last row's 5 cells to make a single row
    last_row = table.rows[8]
    merged_cell = last_row.cells[1].merge(last_row.cells[2])
    for i in range(3, 6):
        merged_cell = merged_cell.merge(last_row.cells[i])

    # Set last row merged column text
    merged_cell.text = "Client asks limit of 7.5mn."
    merged_paragraph = merged_cell.paragraphs[0]
    #merged_paragraph.runs[0].bold = True
    merged_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


    document.add_paragraph()  # One blank line
    document.add_paragraph()  # One blank line

    # --- 3 positives Table ---
    # Create a table with 5 rows (1 for title, 4 for data) and 2 columns
    table = document.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # Merge all cells in the first row for title
    title_cell = table.cell(0, 0).merge(table.cell(0, 1))
    title_cell.text = "3 Positives"
    title_para = title_cell.paragraphs[0]
    title_para.runs[0].bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set background color for header cell
    shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    title_cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    positives = [
        ("Large-Scale Operations", "The company operates on a large scale with multiple branches and facilities across Saudi Arabia."),
        ("Experienced Management.", "The company has over 25 years of experience in the metal and iron industry."),
        ("Strong Sales Growth", "Annual sales increased to SAR 197.56 million YTD, a 24% rise year-on-year."),
        ("Clean Credit History", "The company has maintained a clean credit history, showing no significant past defaults")
    ]

    # Fill in rows 1–4 with the data
    for idx, (title, desc) in enumerate(positives):
        row = table.rows[idx + 1]
        cell1 = row.cells[0]
        cell2 = row.cells[1]

        run = cell1.paragraphs[0].add_run(title)
        run.bold = True
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell2.text = desc
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


    document.add_paragraph()  # One blank line
    document.add_paragraph()  # One blank line

    # --- 3 Negatives Table ---
    # Create a table with 5 rows (1 for title, 4 for data) and 2 columns
    table = document.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # Merge all cells in the first row for title
    title_cell = table.cell(0, 0).merge(table.cell(0, 1))
    title_cell.text = "3 Negatives"
    title_para = title_cell.paragraphs[0]
    title_para.runs[0].bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set background color for header cell
    shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    title_cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    positives = [
        ("High Receivables Concentration", "82% of receivables, totaling SAR 32 million, are from a single client (Al Arabat and Ashghal Factory Co)."),
        ("Leverage Increase", "The leverage ratio jumped from 1.99x in 2023 to 2.93x in 2024."),
        ("Weak Debt Service Coverage", "DSCR is low at 0.41x"),
        ("Past Losses", "The company incurred losses in 2023 due to rising iron prices and the launch of a new rebar plant.")
    ]

    # Fill in rows 1–4 with the data
    for idx, (title, desc) in enumerate(positives):
        row = table.rows[idx + 1]
        cell1 = row.cells[0]
        cell2 = row.cells[1]

        run = cell1.paragraphs[0].add_run(title)
        run.bold = True
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell2.text = desc
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


    document.add_page_break()

    # --- Credit Application Analysis ---
    document.add_heading('Credit Application Analysis', level=1)
    document.add_heading("1. Company's Legal Structure & Management", level=2)

    # === OUTER TABLE ===
    outer_table = document.add_table(rows=2, cols=1)
    outer_table.style = "Table Grid"

    # Add Title
    outer_cell = outer_table.cell(0, 0)
    outer_cell.text = "SHAREHOLDERS / MANAGEMENTS"
    outer_cell.paragraphs[0].runs[0].bold = True
    # Set background color for header cell
    shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    outer_cell._tc.get_or_add_tcPr().append(shading)

    # === INNER TABLE ===
    outer_cell = outer_table.cell(1, 0)
    inner_table = outer_cell.add_table(rows=1, cols=5)
    inner_table.style = "Table Grid"

    # Row 1: SHAREHOLDERS Header Row
    hdr_row = inner_table.rows[0]
    hdr_row.cells[0].text = "SHAREHOLDERS"
    hdr_row.cells[0]._tc.merge(hdr_row.cells[4]._tc)
    hdr_row.cells[0].paragraphs[0].runs[0].bold = True

    # Row 2: Column Headers
    row2 = inner_table.add_row()
    headers = ["Partner Name", "% of ownership", "Email & Mobile#", "Nationality", "Age"]
    for i, title in enumerate(headers):
        cell = row2.cells[i]
        cell.text = title
        cell.paragraphs[0].runs[0].bold = True
        # Set background color for header cell
        shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    # Shareholder Data Rows
    shareholders = [
        ["Anwar Mohammed Alhashari", "40%", "531055777", "Yamen", ""],
        ["Fahad Mohammed Alhabardi", "59%", "503409909", "Saudi", ""],
        ["Ahmed Suliman Alasmri", "1%", "--", "Saudi", ""]
    ]
    for data in shareholders:
        row = inner_table.add_row()
        for i, val in enumerate(data):
            row.cells[i].text = val

    # Row: Authorized Person(s)
    row = inner_table.add_row()
    row.cells[0].text = "Authorized Person(s)"
    row.cells[0].paragraphs[0].runs[0].bold = True
    # Set background color for header cell
    shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
    row.cells[0]._tc.get_or_add_tcPr().append(shading)

    row.cells[1].merge(row.cells[4])
    row.cells[1].text = "Anwar Mohammed Alhashari"

    # Row: MANAGEMENT Header
    row = inner_table.add_row()
    row.cells[0].text = "MANAGEMENT"
    row.cells[0]._tc.merge(row.cells[4]._tc)
    row.cells[0].paragraphs[0].runs[0].bold = True

    # Row: Management Headers
    row = inner_table.add_row()
    mgmt_headers = ["Name", "Title", "Email & Mobile#", "Years of experience", "Years with Co."]
    for i, title in enumerate(mgmt_headers):
        cell = row.cells[i]
        cell.text = title
        cell.paragraphs[0].runs[0].bold = True
        # Set background color for header cell
        shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    # Management Data
    row = inner_table.add_row()
    mgmt_data = ["Anwar Mohammed Alhashari", "GM", "531055777", "--", "--"]
    for i, val in enumerate(mgmt_data):
        row.cells[i].text = val


    document.add_paragraph()  # One blank line

    # 2. Business Overview, Activities and Main Customers
    document.add_heading('2. Business Overview, Activities and Main Customers', level=2)

    # table business overview
    table = document.add_table(rows=2, cols=1)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Business Overview'
    for cell in hdr_cells:
        #cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set background color for header cell
        shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    # Data Row
    data_cells = table.rows[1].cells
    data_cells[0].text = 'Etihad Jeddah Metal Products Factory Co. specializes in scrap collection, iron and steel manufacturing (including rebar and rolled products), and metal trading, with multiple production and logistics facilities across Saudi Arabia.'
    for cell in data_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


    document.add_paragraph()  # One blank line
    document.add_paragraph()  # One blank line

    # Product / Service Table
    table = document.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product / Service '
    hdr_cells[1].text = 'Industry'
    hdr_cells[2].text = 'Target Market (retail/wholesale)'
    hdr_cells[3].text = '% of Annual Revenues'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set background color for header cell
        shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    # Data Row
    data_cells = table.rows[1].cells
    data_cells[0].text = 'Metal Products'
    data_cells[1].text = 'Manufacturing '
    data_cells[2].text = 'Wholesale'
    data_cells[3].text = '100%'
    for cell in data_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Major Customers’ name Table
    document.add_paragraph() # Spacing
    table = document.add_table(rows=4, cols=6)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Major Customers’ name'
    hdr_cells[1].text = 'Country'
    hdr_cells[2].text = 'Total annual sales to Customer,(SAR)'
    hdr_cells[3].text = 'Payment Methods (Bank LC / Transfer/ Cheques)'
    hdr_cells[4].text = 'Payment Terms (days)'
    hdr_cells[5].text = 'Age of Relationship'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set background color for header cell
        shading = parse_xml(r'<w:shd {} w:fill="DEEBF6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    # Data Rows
    data = [
        ['Ajeej Steel Manufacturing Co', 'KSA', '13.69mn', 'Transfer', '180 Day', '--'],
        ['Abdulkarim Alrajhi Steel', 'KSA', '2.27mn', 'Transfer', '180 Day', '--'],
        ['Kema Gulg Trading Co', 'KSA', '22.71mn', 'Transfer', '180 Day', '--']
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Comments
    comments_para = document.add_paragraph('Comments:,Payment terms verified through a bank statement? Yes')
    comments_para.runs[0].bold = True
    comments_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_page_break()

    # 3. Credit Bureau
    document.add_heading('3. Credit Bureau', level=2)

    document.add_paragraph('Company / Branches Credit History', style='Body Text').runs[0].bold = True

    # Report Date
    document.add_paragraph('Report Date,28-05-2025')

    # Financial Institution Name Table
    table = document.add_table(rows=5, cols=5) # Changed from 4 to 5 rows
    table.style = 'Table Grid'

    # Header Row 1
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Financial Institution Name'
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[2].text = 'Total Limit'
    hdr_cells[2].merge(hdr_cells[3])
    hdr_cells[4].text = 'Total Utilization'

    for cell in [hdr_cells[0], hdr_cells[2], hdr_cells[4]]:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header Row 2 (Sub-headers for Limit and Utilization)
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Funded'
    hdr_cells[2].text = 'Non-funded'
    hdr_cells[3].text = 'Funded'
    hdr_cells[4].text = 'Non-funded'
    for cell in hdr_cells[1:]:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data Rows
    data = [
        ['SNB', '20.0mn', '--', '20.0mn', '--'],
        ['Yanal Finance Co', '8.0mn', '--', '6.36mn', '--'],
        ['Total', '28.0mn', '--', '26.36mn', '--']
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i+2].cells # This will now work correctly for all 3 data rows
        cells[0].text = row_data[0]
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j, text in enumerate(row_data[1:]):
            cells[j+1].text = text
            cells[j+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if row_data[0] == 'Total':
            for cell in cells:
                cell.paragraphs[0].runs[0].bold = True

    # Defaults, Bounced Cheques etc.
    document.add_paragraph('Defaults,0,Bounced Cheques,0')
    document.add_paragraph('Past Dues,0,Court Cases,0')
    document.add_paragraph('Comments:,Satisfactory report.')

    document.add_paragraph('List of the company’s facilities as advised by the client:')

    # List of facilities Table
    table = document.add_table(rows=4, cols=5)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Bank Name/ financial Institution'
    hdr_cells[1].text = '"Type of facility (working Capital, Long Term, etc.)"'
    hdr_cells[2].text = 'Total limits'
    hdr_cells[3].text = 'Collateral / Security'
    hdr_cells[4].text = 'Age of Relationship'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data Rows
    data = [
        ['SNB', 'Short Term', '20.0mn', '--', '--'],
        ['Yanal Co', 'Long term', '8.0mn', '--', '--'],
        ['Total Limits', '28.0mn', '--', '--', '--']
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if row_data[0] == 'Total Limits': # Bold 'Total Limits' row
            for cell in cells:
                cell.paragraphs[0].runs[0].bold = True

    document.add_paragraph('Individuals’ SIMAH Report')

    # Individuals’ SIMAH Report (Anawar Alhashri)
    table = document.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    data = [
        ['Individual Name', 'Anawar Alhashri', 'Report Date', '28-05-2025'],
        ['Defaults', '0', 'Bounced Cheques', '0'],
        ['Past Dues', '0', 'Number of Active Trades', '6'],
        ['Total Limits', '848K', 'Total Outstanding', '850K'],
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            # Bold first column items in each row
            if j % 2 == 0:
                cells[j].paragraphs[0].runs[0].bold = True
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph('Total monthly installments (Personal & Car Loans),17K')
    document.add_paragraph('Comments:,Satisfactory.')

    document.add_paragraph('Individuals’ SIMAH Report')

    # Individuals’ SIMAH Report (Fahad Alhabrdi)
    table = document.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    data = [
        ['Individual Name', 'Fahad Alhabrdi ', 'Report Date', '28-05-2025'],
        ['Defaults', '73K', 'Bounced Cheques', '0'],
        ['Past Dues', '0', 'Number of Active Trades', '4'],
        ['Total Limits', '55K', 'Total Outstanding', '55K'],
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            # Bold first column items in each row
            if j % 2 == 0:
                cells[j].paragraphs[0].runs[0].bold = True
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph('Total monthly installments (Personal & Car Loans),-')
    document.add_paragraph('Comments:,"The client has a default with Al Rajhi Bank amounting to SAR 73,000 since 2014, and it remains outstanding."')

    # 4) Financial Analysis (LINK)
    document.add_heading('4) Financial Analysis (LINK)', level=2)
    document.add_paragraph('Is the full year audited? Which is the most recent audited financial?')
    document.add_paragraph('All figures and ratios that include P&L items must be annualized')

    # Financial Analysis Table
    table = document.add_table(rows=13, cols=4)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '' # Empty for first cell
    hdr_cells[1].text = '2024'
    hdr_cells[2].text = '2023'
    hdr_cells[3].text = 'Remarks'
    for cell in hdr_cells[1:]:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data Rows
    data = [
        ['Lendo Facility to Borrower Annual Sales', '0.5%', '0.5%', 'Based on 3.5mn'],
        ['Annual Sales', '197.56mn', '158.92mn', 'In line With VAT Returns'],
        ['Gross Profit ', '7.06mn,GPM (4%)', '-1.46mn,(-1%)', '--'],
        ['Net Profit', '8.77mn,NPM (4%)', '-9.75mn,(-6%)', '"The company incurred losses in 2023 due to rising iron prices, compounded by the opening of its rebar production plant."'],
        ['Cash Conversion Cycle', '37 Days', '27 Days', 'The increase in CCC Due to increase in Receivables Days'],
        ['Sales Growth YoY %', '24%', '-43%', ''],
        ['External Debt / Sales', '13%', '22%', '"The external debt ratio has decreased, as reported by SIMAH, because most of the liabilities are short-term and we are in mid-year."'],
        ['Leverage ratio (Total Liabilities / Net Worth)', '2.93x', '1.99x', 'The leverage increased due to the rise in current liabilities from 49mn in 23 to 101mn in 24'],
        ['Gearing Ratio ', '(External Debt / Net Worth)', '1x', '1x,--'],
        ['Debt Service Coverage Ratio: ', 'EBITDA/(ST Debt + Current portion of long term debt + interest)', '0.41x', '0x,--'],
        ['Current Ratio (Current Asset / Current Liability)', '0.57x', '0.59x', '--'],
        ['Receivables Aging & Concentration ', '31% 0-30 Days,26% 31-60 Days,42% 61-90 Days,1% Above one year', '30% 0-30 Days,51% 31-60 Days,18% 61-90 Days,', ',"The total receivables in 2024 amounted to 32mn, with 82% of them from a single customer (Al Arabat and Ashghal Factory Co)."']
    ]

    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells # Start from 2nd row
        cells[0].text = row_data[0]
        cells[0].paragraphs[0].runs[0].bold = True # Bold the first column for each row
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j, text in enumerate(row_data[1:]):
            cells[j+1].text = text
            cells[j+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 5) Bank Statement Analysis
    document.add_heading('5) Bank Statement Analysis', level=2)

    # Bank Statement Analysis Table
    table = document.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    data = [
        ['Opening Balance', '-'],
        ['Closing balance', '-'],
        ['Total Credit over the last 12 months ', '239.11mn'],
        ['Total Debit over the last 12 months ', '237.67mn'],
        ['"Add other important items like recurring expenses, revenues, etc"', '--']
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].bold = True # Bold the first column

    # 6) Maximum Receivables Financing Required Limit from all FIs:
    document.add_heading('6) Maximum Receivables Financing Required Limit from all FIs:', level=2)

    # Max. Required Limit Table
    table = document.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Latest 12 months'
    hdr_cells[1].paragraphs[0].runs[0].bold = True
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data Rows
    data = [
        ['Revenues (A)', '197.56mn'],
        ['Receivables Days (B)', '61 Days'],
        ['Max. Required Limit (A * B / 365)', '33.01mn']
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].bold = True # Bold first column

    document.add_paragraph('Comments (if any): --')

    # 7) Client Relationship with Lendo (if existing client)
    document.add_heading('7) Client Relationship with Lendo (if existing client)', level=2)

    document.add_paragraph('Details of loans with delayed settlement (PD) during the reviewed period')

    # Loans with delayed settlement Table
    table = document.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Loan#'
    hdr_cells[1].text = 'Loan Amount'
    hdr_cells[2].text = '# of Past Due days'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Empty Data Rows (as per original)
    for i in range(1, 3):
        cells = table.rows[i].cells
        cells[0].text = ''
        cells[1].text = ''
        cells[2].text = ''


    # Save the document
    try:
        document.save(output_filename)
        print(f"Document '{output_filename}' created successfully.")
    except Exception as e:
        print(f"Error saving document: {e}")

# Call the function to create the document
if __name__ == "__main__":
    # Sample data for testing
    summary_data = {
        "companyName": "شركة الأقتصاد الأفتراضي للتجارة",
        "crNumber": "1234567890",
        "simahScore": 789,
        "dpd": "30 days",
        "revenue": "2,000,000 SAR",
        "netProfitMargin": "15%",
        "dscr": "2.1",
        "bouncedCheques": "None",
        "riskRating": "Low",
        "finalRecommendation": "Approve",
        "finalDecision": "Approved"
    }
    create_lendo_credit_file(summary_data)